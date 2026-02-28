"""Load .txt, .docx, and .pdf documents from a directory with source tracking."""

from pathlib import Path
from typing import NamedTuple

from docx import Document as DocxDocument
from PyPDF2 import PdfReader


class Document(NamedTuple):
    """A loaded document with its source filename."""
    text: str
    source: str


def _read_docx(filepath: Path) -> str:
    """Extract all paragraph text from a .docx file."""
    doc = DocxDocument(str(filepath))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_pdf(filepath: Path) -> str:
    """Extract all text from a .pdf file."""
    reader = PdfReader(str(filepath))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def load_documents(data_path: str) -> list[Document]:
    """Read all .txt, .docx, and .pdf files from *data_path*."""
    docs: list[Document] = []
    directory = Path(data_path)
    if not directory.is_dir():
        return docs
    for filepath in sorted(directory.iterdir()):
        suffix = filepath.suffix.lower()
        if suffix == ".txt":
            text = filepath.read_text(encoding="utf-8", errors="ignore").strip()
        elif suffix == ".docx":
            text = _read_docx(filepath).strip()
        elif suffix == ".pdf":
            text = _read_pdf(filepath).strip()
        else:
            continue
        if text:
            docs.append(Document(text=text, source=filepath.name))
    return docs
